from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse, FileResponse
from app.utils.pdf_parser import extract_text_from_pdf
from app.utils.docx_parser import extract_text_from_docx
from app.utils.helpers import write_temp_file, run_blocking
from openai import OpenAI
from docx import Document
import os
import json
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="AI Document Compliance Checker v3")

# Initialize OpenAI client
api_key = os.getenv("OPENAI_API_KEY")
if api_key and api_key.startswith("="):
    api_key = api_key.lstrip("=")

client = OpenAI(api_key=api_key)

@app.post("/analyze_file")
async def analyze_file(file: UploadFile = File(...)):
    """Analyze document for compliance, grammar, and sentiment."""
    try:
        tmp_path = await write_temp_file(file)
        filename = file.filename.lower()

        # Detect file type
        if filename.endswith(".pdf"):
            text = await run_blocking(extract_text_from_pdf, tmp_path)
        elif filename.endswith(".docx"):
            text = await run_blocking(extract_text_from_docx, tmp_path)
        else:
            return JSONResponse(content={"detail": "Unsupported file type"}, status_code=400)

        # Import the analyzer from ai_agent.py
        from app.utils.ai_agent import analyze_text_with_ai
        result = analyze_text_with_ai(text)

        return JSONResponse(content=result, status_code=200)

    except Exception as e:
        return JSONResponse(content={"detail": f"Internal error: {str(e)}"}, status_code=500)



def analyze_text_with_ai(text: str):
    """
    Performs compliance, grammar scoring, and sentiment analysis using GPT.
    """
    prompt = f"""
    You are an AI compliance and linguistic analysis expert.
    Analyze the following text for:
    - Grammar correctness (0–100 score)
    - Tone and sentiment (positive, neutral, negative)
    - Clarity and professionalism
    - Structural and formatting recommendations

    Return your output strictly in JSON format:
    {{
      "summary": "<Brief overview of document>",
      "grammar_score": <number between 0 and 100>,
      "sentiment": "<positive | neutral | negative>",
      "recommendations": ["list of recommendations"],
      "compliance_score": "<percentage>"
    }}

    Text to analyze:
    {text[:5000]}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert in grammar scoring, tone analysis, and compliance evaluation."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )

    result = response.choices[0].message.content

    try:
        clean_json = result.strip().replace("```json", "").replace("```", "")
        return json.loads(clean_json)
    except Exception:
        return {
            "summary": "Error parsing AI response.",
            "grammar_score": 0,
            "sentiment": "unknown",
            "recommendations": [],
            "compliance_score": "N/A",
        }


def correct_text_with_ai(text: str):
    """
    AI correction for tone, grammar, and readability.
    Returns a cleaned and refined version.
    """
    prompt = f"""
    You are a professional editor.
    Correct the following text for:
    - Grammar and punctuation
    - Tone consistency and readability
    - Clarity and professional compliance

    Return only the corrected version — no JSON or explanation.

    Text:
    {text[:7000]}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert copy editor and compliance writer."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
    )

    return response.choices[0].message.content.strip()


@app.post("/correct_file")
async def correct_file(file: UploadFile = File(...)):
    """Correct text and return as DOCX or PDF depending on input type."""
    try:
        # Step 1: Save uploaded file temporarily
        tmp_path = await write_temp_file(file)
        filename = file.filename.lower()

        # Step 2: Detect and extract text
        if filename.endswith(".pdf"):
            text = await run_blocking(extract_text_from_pdf, tmp_path)
        elif filename.endswith(".docx"):
            text = await run_blocking(extract_text_from_docx, tmp_path)
        else:
            return JSONResponse(content={"detail": "Unsupported file type"}, status_code=400)

        # Step 3: AI correction
        corrected_text = correct_text_with_ai(text)

        # Step 4: Create corrected DOCX
        corrected_docx_path = tmp_path.replace(".pdf", "_corrected.docx").replace(".docx", "_corrected.docx")
        doc = Document()
        doc.add_heading("AI Corrected Document", level=1)
        doc.add_paragraph(corrected_text)
        doc.save(corrected_docx_path)

        # Step 5: Try converting DOCX → PDF (fallback if fails)
        corrected_pdf_path = corrected_docx_path.replace(".docx", ".pdf")
        try:
            from docx2pdf import convert
            convert(corrected_docx_path, corrected_pdf_path)
        except Exception as e:
            print("⚠️ PDF conversion failed, falling back to DOCX:", e)
            corrected_pdf_path = None

        # Step 6: Return correct file type
        if corrected_pdf_path and os.path.exists(corrected_pdf_path):
            return FileResponse(
                corrected_pdf_path,
                media_type="application/pdf",
                filename=os.path.basename(corrected_pdf_path),
            )
        else:
            return FileResponse(
                corrected_docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(corrected_docx_path),
            )

    except Exception as e:
        # Step 7: Handle any server-side error
        return JSONResponse(content={"detail": f"Internal error: {str(e)}"}, status_code=500)