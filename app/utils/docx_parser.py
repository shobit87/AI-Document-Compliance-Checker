from docx import Document

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file including paragraphs, tables, headers, and footers.
    """
    doc = Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    # Optional: Add section headers or footer text (if any)
    try:
        for section in doc.sections:
            header = section.header
            footer = section.footer
            if header and header.paragraphs:
                full_text.append("HEADER: " + " ".join(p.text for p in header.paragraphs if p.text.strip()))
            if footer and footer.paragraphs:
                full_text.append("FOOTER: " + " ".join(p.text for p in footer.paragraphs if p.text.strip()))
    except Exception:
        pass  # Some files may not have headers/footers

    return "\n".join(full_text)

